"""
Marketing Campaign Performance Analysis
========================================
Analyzes 200K marketing campaigns across channels, companies, and audience segments.
Generates all figures and the final Word document (.docx).

Usage:
    python analyze.py
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from scipy import stats
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from sklearn.metrics import r2_score, mean_absolute_error
from sklearn.preprocessing import LabelEncoder
import os
import warnings
warnings.filterwarnings('ignore')

# ── Config ──────────────────────────────────────────────────────────────────
DATA_PATH = os.path.join(os.path.dirname(__file__), '..', 'marketing_campaign_dataset.csv')
FIG_DIR = os.path.join(os.path.dirname(__file__), 'figures')
OUTPUT_DOCX = os.path.join(os.path.dirname(__file__), 'Marketing_Campaign_Performance_Analysis.docx')

# Brand palette
PALETTE = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#3B1F2B']
CHANNEL_COLORS = {
    'Google Ads': '#4285F4', 'YouTube': '#FF0000', 'Instagram': '#E1306C',
    'Facebook': '#1877F2', 'Email': '#34A853', 'Website': '#FFC107'
}
sns.set_theme(style='whitegrid', font_scale=1.1, palette=PALETTE)
plt.rcParams['figure.dpi'] = 150
plt.rcParams['savefig.bbox'] = 'tight'
plt.rcParams['font.family'] = 'sans-serif'

# ── Load & Clean ────────────────────────────────────────────────────────────
print("Loading data...")
df = pd.read_csv(DATA_PATH)
df['Acquisition_Cost'] = df['Acquisition_Cost'].str.replace(r'[\$,]', '', regex=True).astype(float)
df['Duration_Days'] = df['Duration'].str.extract(r'(\d+)').astype(int)
df['Date'] = pd.to_datetime(df['Date'])
df['Month'] = df['Date'].dt.to_period('M')
df['Quarter'] = df['Date'].dt.to_period('Q')
df['CTR'] = df['Clicks'] / df['Impressions']
df['CPA'] = df['Acquisition_Cost'] / (df['Clicks'] * df['Conversion_Rate']).replace(0, np.nan)
df['Est_Conversions'] = df['Clicks'] * df['Conversion_Rate']

print(f"Loaded {len(df):,} campaigns | {df['Date'].min().date()} to {df['Date'].max().date()}")

# ── Analysis Functions ──────────────────────────────────────────────────────

def fig_path(name):
    return os.path.join(FIG_DIR, f'{name}.png')


def plot_channel_overview():
    """Fig 1: Channel performance overview (ROI, Conversion Rate, Engagement)"""
    ch = df.groupby('Channel_Used').agg(
        Avg_ROI=('ROI', 'mean'),
        Avg_CR=('Conversion_Rate', 'mean'),
        Avg_Engagement=('Engagement_Score', 'mean'),
        Total_Clicks=('Clicks', 'sum'),
        Campaigns=('Campaign_ID', 'count')
    ).reset_index()

    fig, axes = plt.subplots(1, 3, figsize=(16, 5))
    colors = [CHANNEL_COLORS.get(c, '#999') for c in ch['Channel_Used']]

    # ROI
    bars = axes[0].barh(ch['Channel_Used'], ch['Avg_ROI'], color=colors, edgecolor='white', linewidth=0.5)
    axes[0].set_xlabel('Average ROI')
    axes[0].set_title('Average ROI by Channel', fontweight='bold', fontsize=13)
    for bar, val in zip(bars, ch['Avg_ROI']):
        axes[0].text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2,
                     f'{val:.2f}', va='center', fontsize=10)

    # Conversion Rate
    bars = axes[1].barh(ch['Channel_Used'], ch['Avg_CR'] * 100, color=colors, edgecolor='white', linewidth=0.5)
    axes[1].set_xlabel('Avg Conversion Rate (%)')
    axes[1].set_title('Avg Conversion Rate by Channel', fontweight='bold', fontsize=13)
    for bar, val in zip(bars, ch['Avg_CR']):
        axes[1].text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2,
                     f'{val*100:.1f}%', va='center', fontsize=10)

    # Engagement Score
    bars = axes[2].barh(ch['Channel_Used'], ch['Avg_Engagement'], color=colors, edgecolor='white', linewidth=0.5)
    axes[2].set_xlabel('Avg Engagement Score (1–10)')
    axes[2].set_title('Avg Engagement Score by Channel', fontweight='bold', fontsize=13)
    for bar, val in zip(bars, ch['Avg_Engagement']):
        axes[2].text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2,
                     f'{val:.2f}', va='center', fontsize=10)

    fig.suptitle('Channel Performance Overview', fontsize=16, fontweight='bold', y=1.02)
    plt.tight_layout()
    plt.savefig(fig_path('01_channel_overview'))
    plt.close()
    return ch


def plot_campaign_type_roi():
    """Fig 2: ROI distribution by campaign type"""
    fig, ax = plt.subplots(figsize=(10, 6))
    order = df.groupby('Campaign_Type')['ROI'].median().sort_values(ascending=False).index
    sns.boxplot(data=df, x='Campaign_Type', y='ROI', order=order, palette=PALETTE, ax=ax,
                fliersize=1, linewidth=1.2)
    ax.set_title('ROI Distribution by Campaign Type', fontsize=15, fontweight='bold')
    ax.set_xlabel('Campaign Type', fontsize=12)
    ax.set_ylabel('ROI', fontsize=12)

    medians = df.groupby('Campaign_Type')['ROI'].median()
    for i, ct in enumerate(order):
        ax.text(i, medians[ct] + 0.15, f'{medians[ct]:.2f}', ha='center', fontsize=10, fontweight='bold')

    plt.tight_layout()
    plt.savefig(fig_path('02_campaign_type_roi'))
    plt.close()


def plot_channel_roi_heatmap():
    """Fig 3: Channel x Campaign Type ROI heatmap"""
    pivot = df.pivot_table(values='ROI', index='Channel_Used', columns='Campaign_Type', aggfunc='mean')
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.heatmap(pivot, annot=True, fmt='.2f', cmap='YlOrRd', linewidths=0.5,
                ax=ax, cbar_kws={'label': 'Avg ROI'})
    ax.set_title('Average ROI: Channel × Campaign Type', fontsize=15, fontweight='bold')
    ax.set_xlabel('Campaign Type', fontsize=12)
    ax.set_ylabel('Channel', fontsize=12)
    plt.tight_layout()
    plt.savefig(fig_path('03_channel_campaign_heatmap'))
    plt.close()
    return pivot


def plot_audience_analysis():
    """Fig 4: Audience segment performance"""
    aud = df.groupby('Target_Audience').agg(
        Avg_ROI=('ROI', 'mean'),
        Avg_CR=('Conversion_Rate', 'mean'),
        Avg_CPA=('CPA', 'mean'),
        Total_Conversions=('Est_Conversions', 'sum')
    ).reset_index().sort_values('Avg_ROI', ascending=False)

    fig, axes = plt.subplots(1, 2, figsize=(14, 6))

    # ROI by audience
    bars = axes[0].bar(aud['Target_Audience'], aud['Avg_ROI'], color=PALETTE, edgecolor='white')
    axes[0].set_title('Average ROI by Target Audience', fontsize=14, fontweight='bold')
    axes[0].set_ylabel('Average ROI')
    axes[0].tick_params(axis='x', rotation=15)
    for bar, val in zip(bars, aud['Avg_ROI']):
        axes[0].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.03,
                     f'{val:.2f}', ha='center', fontsize=10, fontweight='bold')

    # Conversions by audience
    bars = axes[1].bar(aud['Target_Audience'], aud['Total_Conversions'], color=PALETTE, edgecolor='white')
    axes[1].set_title('Total Estimated Conversions by Audience', fontsize=14, fontweight='bold')
    axes[1].set_ylabel('Conversions')
    axes[1].tick_params(axis='x', rotation=15)
    axes[1].yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f'{x/1e6:.1f}M'))
    for bar, val in zip(bars, aud['Total_Conversions']):
        axes[1].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 5000,
                     f'{val/1e6:.2f}M', ha='center', fontsize=10, fontweight='bold')

    plt.tight_layout()
    plt.savefig(fig_path('04_audience_analysis'))
    plt.close()
    return aud


def plot_cost_efficiency():
    """Fig 5: Acquisition Cost vs ROI scatter with channel coloring"""
    fig, ax = plt.subplots(figsize=(10, 7))
    for ch, color in CHANNEL_COLORS.items():
        subset = df[df['Channel_Used'] == ch].sample(min(2000, len(df[df['Channel_Used'] == ch])), random_state=42)
        ax.scatter(subset['Acquisition_Cost'], subset['ROI'], c=color, label=ch,
                   alpha=0.35, s=15, edgecolors='none')

    ax.set_xlabel('Acquisition Cost ($)', fontsize=12)
    ax.set_ylabel('ROI', fontsize=12)
    ax.set_title('Cost Efficiency: Acquisition Cost vs. ROI by Channel', fontsize=15, fontweight='bold')
    ax.legend(title='Channel', loc='upper right', framealpha=0.9)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f'${x/1000:.0f}K'))
    plt.tight_layout()
    plt.savefig(fig_path('05_cost_efficiency'))
    plt.close()


def plot_monthly_trends():
    """Fig 6: Monthly trends for key metrics"""
    monthly = df.groupby(df['Date'].dt.to_period('M')).agg(
        Avg_ROI=('ROI', 'mean'),
        Avg_CR=('Conversion_Rate', 'mean'),
        Total_Clicks=('Clicks', 'sum'),
        Avg_Engagement=('Engagement_Score', 'mean')
    ).reset_index()
    monthly['Date'] = monthly['Date'].dt.to_timestamp()

    fig, axes = plt.subplots(2, 2, figsize=(14, 9))

    axes[0,0].plot(monthly['Date'], monthly['Avg_ROI'], color=PALETTE[0], linewidth=2, marker='o', markersize=3)
    axes[0,0].set_title('Monthly Avg ROI', fontweight='bold')
    axes[0,0].set_ylabel('ROI')

    axes[0,1].plot(monthly['Date'], monthly['Avg_CR']*100, color=PALETTE[1], linewidth=2, marker='o', markersize=3)
    axes[0,1].set_title('Monthly Avg Conversion Rate', fontweight='bold')
    axes[0,1].set_ylabel('Conversion Rate (%)')

    axes[1,0].bar(monthly['Date'], monthly['Total_Clicks'], color=PALETTE[2], width=20)
    axes[1,0].set_title('Monthly Total Clicks', fontweight='bold')
    axes[1,0].set_ylabel('Clicks')
    axes[1,0].yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f'{x/1e6:.1f}M'))

    axes[1,1].plot(monthly['Date'], monthly['Avg_Engagement'], color=PALETTE[3], linewidth=2, marker='o', markersize=3)
    axes[1,1].set_title('Monthly Avg Engagement Score', fontweight='bold')
    axes[1,1].set_ylabel('Score (1–10)')

    for ax in axes.flat:
        ax.tick_params(axis='x', rotation=45)
    fig.suptitle('Campaign Performance Trends (Jan 2021 – Dec 2021)', fontsize=16, fontweight='bold', y=1.01)
    plt.tight_layout()
    plt.savefig(fig_path('06_monthly_trends'))
    plt.close()


def plot_customer_segment():
    """Fig 7: Customer segment deep-dive"""
    seg = df.groupby('Customer_Segment').agg(
        Avg_ROI=('ROI', 'mean'),
        Avg_CR=('Conversion_Rate', 'mean'),
        Avg_Engagement=('Engagement_Score', 'mean'),
        Avg_CPA=('CPA', 'mean'),
        Campaigns=('Campaign_ID', 'count')
    ).reset_index().sort_values('Avg_ROI', ascending=False)

    fig, axes = plt.subplots(1, 3, figsize=(16, 5.5))

    # ROI
    axes[0].barh(seg['Customer_Segment'], seg['Avg_ROI'], color=PALETTE, edgecolor='white')
    axes[0].set_xlabel('Avg ROI')
    axes[0].set_title('ROI by Segment', fontweight='bold')

    # CR
    axes[1].barh(seg['Customer_Segment'], seg['Avg_CR']*100, color=PALETTE, edgecolor='white')
    axes[1].set_xlabel('Avg Conversion Rate (%)')
    axes[1].set_title('Conversion Rate by Segment', fontweight='bold')

    # Engagement
    axes[2].barh(seg['Customer_Segment'], seg['Avg_Engagement'], color=PALETTE, edgecolor='white')
    axes[2].set_xlabel('Avg Engagement Score')
    axes[2].set_title('Engagement by Segment', fontweight='bold')

    fig.suptitle('Customer Segment Performance', fontsize=16, fontweight='bold', y=1.02)
    plt.tight_layout()
    plt.savefig(fig_path('07_customer_segments'))
    plt.close()
    return seg


def plot_company_comparison():
    """Fig 8: Company performance comparison"""
    comp = df.groupby('Company').agg(
        Avg_ROI=('ROI', 'mean'),
        Avg_CR=('Conversion_Rate', 'mean'),
        Total_Spend=('Acquisition_Cost', 'sum'),
        Total_Conversions=('Est_Conversions', 'sum'),
        Avg_Engagement=('Engagement_Score', 'mean')
    ).reset_index().sort_values('Avg_ROI', ascending=False)

    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    x = range(len(comp))
    w = 0.35

    axes[0].bar([i - w/2 for i in x], comp['Avg_ROI'], w, label='Avg ROI', color=PALETTE[0])
    axes[0].bar([i + w/2 for i in x], comp['Avg_CR']*100, w, label='Avg CR (%)', color=PALETTE[1])
    axes[0].set_xticks(list(x))
    axes[0].set_xticklabels(comp['Company'], rotation=15, ha='right', fontsize=9)
    axes[0].set_title('ROI & Conversion Rate by Company', fontweight='bold')
    axes[0].legend()

    axes[1].bar(comp['Company'], comp['Total_Spend'] / 1e9, color=PALETTE[2], edgecolor='white')
    axes[1].set_title('Total Campaign Spend by Company', fontweight='bold')
    axes[1].set_ylabel('Total Spend ($B)')
    axes[1].tick_params(axis='x', rotation=15)

    plt.tight_layout()
    plt.savefig(fig_path('08_company_comparison'))
    plt.close()
    return comp


def plot_duration_impact():
    """Fig 9: Campaign duration impact on performance"""
    dur = df.groupby('Duration_Days').agg(
        Avg_ROI=('ROI', 'mean'),
        Avg_CR=('Conversion_Rate', 'mean'),
        Avg_Engagement=('Engagement_Score', 'mean')
    ).reset_index().sort_values('Duration_Days')

    fig, ax = plt.subplots(figsize=(9, 6))
    x = range(len(dur))
    w = 0.25
    ax.bar([i - w for i in x], dur['Avg_ROI'], w, label='Avg ROI', color=PALETTE[0])
    ax.bar(list(x), dur['Avg_CR'] * 100, w, label='Avg CR (%)', color=PALETTE[1])
    ax.bar([i + w for i in x], dur['Avg_Engagement'], w, label='Avg Engagement', color=PALETTE[2])
    ax.set_xticks(list(x))
    ax.set_xticklabels([f'{d} days' for d in dur['Duration_Days']])
    ax.set_title('Campaign Duration Impact on Key Metrics', fontsize=14, fontweight='bold')
    ax.legend()
    plt.tight_layout()
    plt.savefig(fig_path('09_duration_impact'))
    plt.close()
    return dur


def plot_location_performance():
    """Fig 10: Geographic performance"""
    loc = df.groupby('Location').agg(
        Avg_ROI=('ROI', 'mean'),
        Total_Clicks=('Clicks', 'sum'),
        Avg_Engagement=('Engagement_Score', 'mean'),
        Avg_CPA=('CPA', 'mean')
    ).reset_index().sort_values('Avg_ROI', ascending=False)

    fig, axes = plt.subplots(1, 2, figsize=(13, 5.5))
    axes[0].bar(loc['Location'], loc['Avg_ROI'], color=PALETTE, edgecolor='white')
    axes[0].set_title('Avg ROI by Location', fontweight='bold')
    axes[0].set_ylabel('ROI')
    for i, (_, row) in enumerate(loc.iterrows()):
        axes[0].text(i, row['Avg_ROI'] + 0.03, f"{row['Avg_ROI']:.2f}", ha='center', fontsize=10, fontweight='bold')

    axes[1].bar(loc['Location'], loc['Total_Clicks'] / 1e6, color=PALETTE, edgecolor='white')
    axes[1].set_title('Total Clicks by Location (Millions)', fontweight='bold')
    axes[1].set_ylabel('Clicks (M)')

    plt.tight_layout()
    plt.savefig(fig_path('10_location_performance'))
    plt.close()
    return loc


def compute_social_media_deep_dive():
    """Deep-dive into social media channel (Instagram, Facebook, YouTube)"""
    social = df[df['Channel_Used'].isin(['Instagram', 'Facebook', 'YouTube'])]
    summary = social.groupby('Channel_Used').agg(
        Campaigns=('Campaign_ID', 'count'),
        Avg_ROI=('ROI', 'mean'),
        Avg_CR=('Conversion_Rate', 'mean'),
        Avg_CTR=('CTR', 'mean'),
        Avg_Engagement=('Engagement_Score', 'mean'),
        Avg_CPA=('CPA', 'mean'),
        Avg_Acq_Cost=('Acquisition_Cost', 'mean'),
        Total_Clicks=('Clicks', 'sum'),
        Total_Impressions=('Impressions', 'sum'),
        Total_Conversions=('Est_Conversions', 'sum')
    ).reset_index()

    fig, axes = plt.subplots(2, 2, figsize=(13, 10))
    sm_colors = ['#E1306C', '#1877F2', '#FF0000']
    channels = ['Instagram', 'Facebook', 'YouTube']

    # ROI comparison
    vals = [summary[summary['Channel_Used']==c]['Avg_ROI'].values[0] for c in channels]
    bars = axes[0,0].bar(channels, vals, color=sm_colors, edgecolor='white')
    axes[0,0].set_title('Average ROI', fontweight='bold', fontsize=13)
    for bar, val in zip(bars, vals):
        axes[0,0].text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.03,
                       f'{val:.2f}', ha='center', fontweight='bold')

    # CTR comparison
    vals = [summary[summary['Channel_Used']==c]['Avg_CTR'].values[0]*100 for c in channels]
    bars = axes[0,1].bar(channels, vals, color=sm_colors, edgecolor='white')
    axes[0,1].set_title('Average CTR (%)', fontweight='bold', fontsize=13)
    for bar, val in zip(bars, vals):
        axes[0,1].text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.1,
                       f'{val:.1f}%', ha='center', fontweight='bold')

    # Engagement
    vals = [summary[summary['Channel_Used']==c]['Avg_Engagement'].values[0] for c in channels]
    bars = axes[1,0].bar(channels, vals, color=sm_colors, edgecolor='white')
    axes[1,0].set_title('Average Engagement Score', fontweight='bold', fontsize=13)
    for bar, val in zip(bars, vals):
        axes[1,0].text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.05,
                       f'{val:.2f}', ha='center', fontweight='bold')

    # Total Conversions
    vals = [summary[summary['Channel_Used']==c]['Total_Conversions'].values[0] for c in channels]
    bars = axes[1,1].bar(channels, vals, color=sm_colors, edgecolor='white')
    axes[1,1].set_title('Total Estimated Conversions', fontweight='bold', fontsize=13)
    axes[1,1].yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f'{x/1e6:.1f}M'))
    for bar, val in zip(bars, vals):
        axes[1,1].text(bar.get_x()+bar.get_width()/2, bar.get_height()+5000,
                       f'{val/1e6:.2f}M', ha='center', fontweight='bold')

    fig.suptitle('Social Media Channel Deep-Dive', fontsize=16, fontweight='bold', y=1.01)
    plt.tight_layout()
    plt.savefig(fig_path('11_social_media_deep_dive'))
    plt.close()
    return summary


# ── NEW: Statistical Tests ──────────────────────────────────────────────────

def run_anova_tests():
    """Run one-way ANOVA for ROI across key categorical variables."""
    results = {}

    # Channel ANOVA
    groups = [g['ROI'].values for _, g in df.groupby('Channel_Used')]
    f_stat, p_val = stats.f_oneway(*groups)
    results['Channel_Used'] = {'F': f_stat, 'p': p_val, 'significant': p_val < 0.05}

    # Campaign Type ANOVA
    groups = [g['ROI'].values for _, g in df.groupby('Campaign_Type')]
    f_stat, p_val = stats.f_oneway(*groups)
    results['Campaign_Type'] = {'F': f_stat, 'p': p_val, 'significant': p_val < 0.05}

    # Location ANOVA
    groups = [g['ROI'].values for _, g in df.groupby('Location')]
    f_stat, p_val = stats.f_oneway(*groups)
    results['Location'] = {'F': f_stat, 'p': p_val, 'significant': p_val < 0.05}

    # Customer Segment ANOVA
    groups = [g['ROI'].values for _, g in df.groupby('Customer_Segment')]
    f_stat, p_val = stats.f_oneway(*groups)
    results['Customer_Segment'] = {'F': f_stat, 'p': p_val, 'significant': p_val < 0.05}

    # Target Audience ANOVA
    groups = [g['ROI'].values for _, g in df.groupby('Target_Audience')]
    f_stat, p_val = stats.f_oneway(*groups)
    results['Target_Audience'] = {'F': f_stat, 'p': p_val, 'significant': p_val < 0.05}

    # Duration ANOVA
    groups = [g['ROI'].values for _, g in df.groupby('Duration_Days')]
    f_stat, p_val = stats.f_oneway(*groups)
    results['Duration'] = {'F': f_stat, 'p': p_val, 'significant': p_val < 0.05}

    return results


def run_predictive_model():
    """Train a Random Forest to predict ROI and extract feature importances."""
    # Encode categorical variables
    feature_cols = ['Campaign_Type', 'Channel_Used', 'Target_Audience', 'Customer_Segment',
                    'Location', 'Language', 'Duration_Days', 'Clicks', 'Impressions',
                    'Engagement_Score', 'Acquisition_Cost', 'Conversion_Rate']

    df_model = df[feature_cols + ['ROI']].copy()
    encoders = {}
    cat_cols = ['Campaign_Type', 'Channel_Used', 'Target_Audience', 'Customer_Segment',
                'Location', 'Language']
    for col in cat_cols:
        le = LabelEncoder()
        df_model[col] = le.fit_transform(df_model[col])
        encoders[col] = le

    X = df_model[feature_cols]
    y = df_model['ROI']
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    rf = RandomForestRegressor(n_estimators=200, max_depth=10, random_state=42, n_jobs=-1)
    rf.fit(X_train, y_train)

    y_pred = rf.predict(X_test)
    r2 = r2_score(y_test, y_pred)
    mae = mean_absolute_error(y_test, y_pred)

    # Feature importances
    importances = pd.DataFrame({
        'Feature': feature_cols,
        'Importance': rf.feature_importances_
    }).sort_values('Importance', ascending=False)

    # Plot feature importances
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = [PALETTE[i % len(PALETTE)] for i in range(len(importances))]
    bars = ax.barh(importances['Feature'][::-1], importances['Importance'][::-1],
                   color=colors[::-1], edgecolor='white')
    ax.set_xlabel('Feature Importance', fontsize=12)
    ax.set_title('Random Forest: Feature Importance for Predicting ROI',
                 fontsize=14, fontweight='bold')
    for bar, val in zip(bars, importances['Importance'][::-1]):
        ax.text(bar.get_width() + 0.005, bar.get_y() + bar.get_height()/2,
                f'{val:.3f}', va='center', fontsize=9)
    plt.tight_layout()
    plt.savefig(fig_path('12_feature_importance'))
    plt.close()

    # Plot predicted vs actual
    fig, ax = plt.subplots(figsize=(8, 8))
    sample_idx = np.random.RandomState(42).choice(len(y_test), 3000, replace=False)
    ax.scatter(y_test.iloc[sample_idx], y_pred[sample_idx], alpha=0.2, s=10, c=PALETTE[0])
    ax.plot([2, 8], [2, 8], 'r--', linewidth=2, label='Perfect Prediction')
    ax.set_xlabel('Actual ROI', fontsize=12)
    ax.set_ylabel('Predicted ROI', fontsize=12)
    ax.set_title(f'Predicted vs. Actual ROI (R² = {r2:.4f})', fontsize=14, fontweight='bold')
    ax.legend(fontsize=11)
    plt.tight_layout()
    plt.savefig(fig_path('13_predicted_vs_actual'))
    plt.close()

    return {'r2': r2, 'mae': mae, 'importances': importances}


def plot_correlation_matrix():
    """Correlation heatmap of numeric variables."""
    numeric_cols = ['Conversion_Rate', 'Acquisition_Cost', 'ROI', 'Clicks',
                    'Impressions', 'Engagement_Score', 'Duration_Days', 'CTR', 'CPA', 'Est_Conversions']
    corr = df[numeric_cols].corr()

    fig, ax = plt.subplots(figsize=(11, 9))
    mask = np.triu(np.ones_like(corr, dtype=bool))
    sns.heatmap(corr, mask=mask, annot=True, fmt='.2f', cmap='RdBu_r', center=0,
                linewidths=0.5, ax=ax, vmin=-1, vmax=1,
                cbar_kws={'label': 'Correlation Coefficient'})
    ax.set_title('Correlation Matrix of Campaign Metrics', fontsize=15, fontweight='bold')
    plt.tight_layout()
    plt.savefig(fig_path('14_correlation_matrix'))
    plt.close()
    return corr


# ── Generate all figures ────────────────────────────────────────────────────
print("Generating figures...")
channel_stats = plot_channel_overview()
plot_campaign_type_roi()
heatmap_data = plot_channel_roi_heatmap()
audience_stats = plot_audience_analysis()
plot_cost_efficiency()
plot_monthly_trends()
segment_stats = plot_customer_segment()
company_stats = plot_company_comparison()
duration_stats = plot_duration_impact()
location_stats = plot_location_performance()
social_stats = compute_social_media_deep_dive()

print("Running statistical tests...")
anova_results = run_anova_tests()
print("Training predictive model...")
model_results = run_predictive_model()
print("Generating correlation matrix...")
corr_matrix = plot_correlation_matrix()

# ── Compute summary statistics for the report ──────────────────────────────
total_campaigns = len(df)
total_spend = df['Acquisition_Cost'].sum()
overall_roi = df['ROI'].mean()
overall_cr = df['Conversion_Rate'].mean()
total_clicks = df['Clicks'].sum()
total_impressions = df['Impressions'].sum()
total_conversions = df['Est_Conversions'].sum()
best_channel = channel_stats.sort_values('Avg_ROI', ascending=False).iloc[0]
best_campaign_type = df.groupby('Campaign_Type')['ROI'].mean().idxmax()
best_audience = audience_stats.iloc[0]

# ── Build Word Document ─────────────────────────────────────────────────────
print("Building Word document...")
doc = Document()

# -- Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_figure(fig_name, caption, width=Inches(6.2)):
    doc.add_picture(fig_path(fig_name), width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_kpi_table(data_dict, title=None):
    """Add a styled KPI summary table."""
    if title:
        p = doc.add_paragraph(title)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
    table = doc.add_table(rows=1, cols=len(data_dict))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, (key, val) in enumerate(data_dict.items()):
        cell = hdr.cells[i]
        # KPI name
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(key)
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # KPI value
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(str(val))
        run2.bold = True
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Style table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): 'DDDDDD'
        })
        borders.append(el)
    tblPr.append(borders)
    doc.add_paragraph()  # spacer

def add_data_table(headers, rows):
    """Add a professional data table."""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '1A1A2E'
        })
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if r % 2 == 1:
                shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5'
                })
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()  # spacer


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Marketing Campaign\nPerformance Analysis')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Data-Driven Evaluation of Multi-Channel Marketing Effectiveness\nAcross Audience Segments, Platforms, and Campaign Strategies')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Rebecca Wu\nApril 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run(f'Dataset: 200,000 campaigns | 5 companies | 6 channels | 5 markets')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Introduction & Objectives',
    '3. Data Overview & Methodology',
    '4. Channel Performance Analysis',
    '5. Campaign Type Effectiveness',
    '6. Social Media Deep-Dive: Instagram vs. Facebook vs. YouTube',
    '7. Audience & Customer Segment Analysis',
    '8. Geographic & Demographic Insights',
    '9. Cost Efficiency & Budget Optimization',
    '10. Temporal Trends & Seasonality',
    '11. Company Benchmarking',
    '12. Statistical Significance Testing (ANOVA)',
    '13. Predictive Modeling & Feature Importance',
    '14. Correlation Analysis',
    '15. Key Findings & Strategic Recommendations',
    '16. Limitations & Future Research',
    '17. Appendix'
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('1. Executive Summary', level=1)

add_body(
    f'This report presents a comprehensive analysis of {total_campaigns:,} marketing campaigns executed '
    f'by five companies across six marketing channels throughout 2021. The analysis evaluates campaign '
    f'performance through multiple lenses—return on investment (ROI), conversion rates, engagement scores, '
    f'click-through rates, and cost efficiency—to identify the highest-performing channels, audience '
    f'segments, and campaign strategies.'
)

add_kpi_table({
    'Total Campaigns': f'{total_campaigns:,}',
    'Total Spend': f'${total_spend/1e9:.2f}B',
    'Avg ROI': f'{overall_roi:.2f}x',
    'Avg Conversion Rate': f'{overall_cr*100:.1f}%',
    'Total Conversions': f'{total_conversions/1e6:.1f}M'
})

add_body('Key findings:')
findings = [
    f'All six channels demonstrate comparable average ROI (~{overall_roi:.1f}x), indicating a well-diversified marketing portfolio with no single point of failure.',
    f'Social media platforms (Instagram, Facebook, YouTube) collectively account for approximately half of all campaigns and deliver engagement scores averaging {df[df["Channel_Used"].isin(["Instagram","Facebook","YouTube"])]["Engagement_Score"].mean():.1f}/10.',
    f'Campaign duration shows minimal impact on ROI, suggesting that budget allocation efficiency matters more than campaign length.',
    f'Customer segments show similar performance profiles, with engagement scores ranging from {segment_stats["Avg_Engagement"].min():.1f} to {segment_stats["Avg_Engagement"].max():.1f}, indicating consistent messaging effectiveness across verticals.',
    f'Total marketing spend of ${total_spend/1e9:.2f}B generated an estimated {total_conversions/1e6:.1f}M conversions across all channels.'
]
for f in findings:
    p = doc.add_paragraph(f, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. INTRODUCTION & OBJECTIVES
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('2. Introduction & Objectives', level=1)

add_body(
    'In today\'s hyper-competitive digital landscape, marketing teams face mounting pressure to demonstrate '
    'measurable returns on every dollar spent. With budgets distributed across email, search, display, '
    'social media, and influencer channels, understanding which combinations of platform, audience, and '
    'campaign type deliver the strongest outcomes is critical for strategic planning.'
)

add_body(
    'This analysis was conducted to answer five core business questions:'
)

objectives = [
    'Channel Effectiveness: Which marketing channels deliver the highest ROI, and how do they compare on engagement and conversion metrics?',
    'Social Media Benchmarking: Among Instagram, Facebook, and YouTube, which platform offers the best cost-adjusted performance for different campaign types?',
    'Audience Optimization: Which target demographics and customer segments respond most favorably to marketing efforts, and how should targeting evolve?',
    'Budget Allocation: Where are acquisition costs highest relative to returns, and how can spend be reallocated for maximum efficiency?',
    'Strategic Planning: What campaign durations, geographic markets, and temporal patterns should inform the 2022 marketing strategy?'
]
for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(f'{obj}')
    p.style = 'List Number'
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. DATA OVERVIEW & METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('3. Data Overview & Methodology', level=1)

add_heading_styled('3.1 Dataset Description', level=2)
add_body(
    f'The dataset comprises {total_campaigns:,} individual campaign records spanning January 1, 2021 '
    f'through December 31, 2021. Each record captures 16 attributes describing campaign configuration, '
    f'performance metrics, and contextual metadata.'
)

add_data_table(
    ['Variable', 'Type', 'Description', 'Example Values'],
    [
        ['Campaign_ID', 'Integer', 'Unique campaign identifier', '1 – 200,000'],
        ['Company', 'Categorical', 'Advertising company', 'TechCorp, NexGen Systems, etc.'],
        ['Campaign_Type', 'Categorical', 'Marketing strategy type', 'Email, Social Media, Search, Display, Influencer'],
        ['Target_Audience', 'Categorical', 'Demographic targeting', 'Men 18-24, Women 25-34, All Ages'],
        ['Duration', 'Ordinal', 'Campaign run length', '15, 30, 45, or 60 days'],
        ['Channel_Used', 'Categorical', 'Distribution platform', 'Google Ads, YouTube, Instagram, Facebook, Email, Website'],
        ['Conversion_Rate', 'Continuous', 'Action completion rate', '0.01 – 0.15'],
        ['Acquisition_Cost', 'Continuous', 'Customer acquisition spend', '$5,000 – $20,000'],
        ['ROI', 'Continuous', 'Return on investment', '2.00 – 8.00'],
        ['Location', 'Categorical', 'Market geography', 'New York, Los Angeles, Chicago, Houston, Miami'],
        ['Clicks', 'Integer', 'User click count', '100 – 1,000'],
        ['Impressions', 'Integer', 'Total ad views', '1,000 – 10,000'],
        ['Engagement_Score', 'Integer', 'Engagement rating (1-10)', '1 – 10'],
        ['Customer_Segment', 'Categorical', 'Behavioral segment', 'Tech Enthusiasts, Fashionistas, etc.'],
        ['Date', 'Date', 'Campaign launch date', '2021-01-01 – 2021-12-31'],
    ]
)

add_heading_styled('3.2 Methodology', level=2)
add_body(
    'The analysis employs descriptive statistics, comparative benchmarking, and visual analytics to evaluate '
    'campaign performance. Key derived metrics include:'
)
derived = [
    'Click-Through Rate (CTR) = Clicks / Impressions — measures ad resonance and creative effectiveness.',
    'Cost Per Acquisition (CPA) = Acquisition Cost / Estimated Conversions — quantifies the unit economics of customer acquisition.',
    'Estimated Conversions = Clicks x Conversion Rate — approximates the total actions driven by each campaign.',
]
for d in derived:
    doc.add_paragraph(d, style='List Bullet')

add_body(
    'All analyses were conducted using Python with the following libraries: pandas for data manipulation, '
    'matplotlib and seaborn for visualization, scipy for statistical hypothesis testing (ANOVA), '
    'scikit-learn for machine learning (Random Forest), and NumPy for numerical computing. '
    'Visualizations were rendered at publication quality. Statistical comparisons use group-level aggregates across the full 200,000-record dataset.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. CHANNEL PERFORMANCE ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('4. Channel Performance Analysis', level=1)

add_body(
    'Understanding channel-level performance is foundational to marketing strategy. This section evaluates '
    'all six channels across ROI, conversion rate, and engagement metrics.'
)

add_figure('01_channel_overview', 'Figure 1: Channel Performance Overview — ROI, Conversion Rate, and Engagement Score across all six marketing channels.')

ch_sorted = channel_stats.sort_values('Avg_ROI', ascending=False)
add_data_table(
    ['Channel', 'Campaigns', 'Avg ROI', 'Avg Conversion Rate', 'Avg Engagement', 'Total Clicks'],
    [[row['Channel_Used'], f"{row['Campaigns']:,}", f"{row['Avg_ROI']:.2f}",
      f"{row['Avg_CR']*100:.1f}%", f"{row['Avg_Engagement']:.2f}", f"{row['Total_Clicks']:,}"]
     for _, row in ch_sorted.iterrows()]
)

add_body(
    f'The analysis reveals a notably balanced channel portfolio. Average ROI ranges narrowly across channels, '
    f'indicating that the marketing teams have achieved effective optimization across platforms. '
    f'This parity suggests mature campaign management practices where underperforming channels have already '
    f'been iterated upon or deprioritized.'
)

add_body(
    'However, the similarity in averages masks important distributional differences. '
    'When examined by campaign type interaction (Section 5), certain channel-campaign combinations '
    'significantly outperform others, suggesting that the "best" channel depends heavily on the campaign objective.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. CAMPAIGN TYPE EFFECTIVENESS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('5. Campaign Type Effectiveness', level=1)

add_body(
    'Five distinct campaign types were deployed across the portfolio: Email, Social Media, Influencer, '
    'Display, and Search. Each serves different funnel stages and audience engagement models.'
)

add_figure('02_campaign_type_roi', 'Figure 2: ROI distribution by campaign type. Box plots show median, interquartile range, and outliers.')

ct_stats = df.groupby('Campaign_Type').agg(
    Avg_ROI=('ROI', 'mean'), Median_ROI=('ROI', 'median'),
    Avg_CR=('Conversion_Rate', 'mean'), Avg_Engagement=('Engagement_Score', 'mean')
).reset_index().sort_values('Avg_ROI', ascending=False)

add_data_table(
    ['Campaign Type', 'Avg ROI', 'Median ROI', 'Avg Conversion Rate', 'Avg Engagement'],
    [[row['Campaign_Type'], f"{row['Avg_ROI']:.2f}", f"{row['Median_ROI']:.2f}",
      f"{row['Avg_CR']*100:.1f}%", f"{row['Avg_Engagement']:.2f}"]
     for _, row in ct_stats.iterrows()]
)

add_heading_styled('5.1 Channel × Campaign Type Interaction', level=2)
add_body(
    'The heatmap below reveals which channel-campaign combinations produce the highest returns. '
    'This interaction analysis is critical for tactical budget allocation.'
)
add_figure('03_channel_campaign_heatmap', 'Figure 3: Average ROI by Channel and Campaign Type interaction. Darker cells indicate higher returns.')

add_body(
    'While the overall averages are similar, the interaction heatmap identifies specific pockets of '
    'outperformance. Marketing teams should use this matrix to guide channel selection based on the '
    'campaign type being deployed.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. SOCIAL MEDIA DEEP-DIVE
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('6. Social Media Deep-Dive: Instagram vs. Facebook vs. YouTube', level=1)

add_body(
    'Social media platforms represent a critical component of the modern marketing mix. This section '
    'provides a head-to-head comparison of Instagram, Facebook, and YouTube across four key performance dimensions.'
)

add_figure('11_social_media_deep_dive',
    'Figure 4: Social media platform comparison across ROI, CTR, Engagement, and Total Conversions.')

sm = social_stats.sort_values('Avg_ROI', ascending=False)
add_data_table(
    ['Platform', 'Campaigns', 'Avg ROI', 'Avg CTR', 'Avg CR', 'Avg Engagement', 'Avg CPA', 'Total Conversions'],
    [[row['Channel_Used'], f"{row['Campaigns']:,}", f"{row['Avg_ROI']:.2f}",
      f"{row['Avg_CTR']*100:.1f}%", f"{row['Avg_CR']*100:.1f}%",
      f"{row['Avg_Engagement']:.2f}", f"${row['Avg_CPA']:,.0f}",
      f"{row['Total_Conversions']/1e6:.2f}M"]
     for _, row in sm.iterrows()]
)

add_heading_styled('6.1 Platform-Specific Insights', level=2)

# Instagram
add_heading_styled('Instagram', level=3)
ig = social_stats[social_stats['Channel_Used'] == 'Instagram'].iloc[0]
add_body(
    f'Instagram processed {ig["Campaigns"]:,} campaigns with an average ROI of {ig["Avg_ROI"]:.2f}x '
    f'and an engagement score of {ig["Avg_Engagement"]:.2f}/10. With its visual-first format, Instagram '
    f'excels at brand storytelling and product discovery, making it particularly effective for Fashionistas '
    f'and lifestyle-oriented customer segments.'
)

# Facebook
add_heading_styled('Facebook', level=3)
fb = social_stats[social_stats['Channel_Used'] == 'Facebook'].iloc[0]
add_body(
    f'Facebook delivered {fb["Campaigns"]:,} campaigns with an average ROI of {fb["Avg_ROI"]:.2f}x. '
    f'Its strength lies in precise demographic targeting and broad reach across age groups. '
    f'Facebook\'s advanced audience tools make it suitable for both awareness and conversion-focused campaigns.'
)

# YouTube
add_heading_styled('YouTube', level=3)
yt = social_stats[social_stats['Channel_Used'] == 'YouTube'].iloc[0]
add_body(
    f'YouTube ran {yt["Campaigns"]:,} campaigns with an average ROI of {yt["Avg_ROI"]:.2f}x '
    f'and generated {yt["Total_Conversions"]/1e6:.2f}M estimated conversions. Video content on YouTube '
    f'offers longer engagement windows and is particularly effective for product demonstrations, '
    f'tutorials, and influencer collaborations.'
)

add_heading_styled('6.2 Social Media Recommendations', level=2)
recs = [
    'Allocate social budgets proportionally based on campaign objectives: use Instagram for brand awareness and visual storytelling, Facebook for precision targeting, and YouTube for long-form educational content.',
    'Test cross-platform sequential campaigns where awareness is built on YouTube/Instagram and conversions are captured through Facebook retargeting.',
    'Monitor platform-specific engagement trends quarterly, as algorithm changes can shift performance significantly.',
]
for r in recs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. AUDIENCE & CUSTOMER SEGMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('7. Audience & Customer Segment Analysis', level=1)

add_heading_styled('7.1 Target Audience Demographics', level=2)
add_body(
    'Campaigns targeted five distinct demographic groups. Understanding differential response rates '
    'enables more efficient media buying and creative development.'
)
add_figure('04_audience_analysis', 'Figure 5: Average ROI and Total Conversions by Target Audience demographic.')

add_heading_styled('7.2 Customer Segment Performance', level=2)
add_body(
    'Beyond demographics, campaigns were mapped to five behavioral segments. These psychographic '
    'profiles influence content strategy and channel selection.'
)
add_figure('07_customer_segments', 'Figure 6: Customer segment performance across ROI, Conversion Rate, and Engagement Score.')

seg_sorted = segment_stats.sort_values('Avg_ROI', ascending=False)
add_data_table(
    ['Customer Segment', 'Campaigns', 'Avg ROI', 'Avg CR', 'Avg Engagement'],
    [[row['Customer_Segment'], f"{row['Campaigns']:,}", f"{row['Avg_ROI']:.2f}",
      f"{row['Avg_CR']*100:.1f}%", f"{row['Avg_Engagement']:.2f}"]
     for _, row in seg_sorted.iterrows()]
)

add_body(
    'Customer segments demonstrate broadly similar performance metrics, suggesting that the current '
    'creative and targeting strategies effectively resonate across behavioral profiles. However, even '
    'marginal differences in conversion rate at this campaign volume translate to thousands of additional conversions, '
    'warranting segment-specific creative testing.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. GEOGRAPHIC & DEMOGRAPHIC INSIGHTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('8. Geographic & Demographic Insights', level=1)

add_body(
    'Campaigns were distributed across five major U.S. metropolitan markets. Geographic performance '
    'differences can inform regional budget allocation and localization strategies.'
)

add_figure('10_location_performance', 'Figure 7: Average ROI and Total Clicks by geographic market.')

loc_sorted = location_stats.sort_values('Avg_ROI', ascending=False)
add_data_table(
    ['Location', 'Avg ROI', 'Total Clicks', 'Avg Engagement', 'Avg CPA'],
    [[row['Location'], f"{row['Avg_ROI']:.2f}", f"{row['Total_Clicks']:,}",
      f"{row['Avg_Engagement']:.2f}", f"${row['Avg_CPA']:,.0f}"]
     for _, row in loc_sorted.iterrows()]
)

add_body(
    'Geographic performance remains relatively uniform, suggesting national-level campaign strategies '
    'are well-calibrated. For markets showing marginally lower ROI, localized creative testing and '
    'channel mix adjustments are recommended before reallocating significant budget.'
)

# Language analysis
lang_stats = df.groupby('Language').agg(
    Avg_ROI=('ROI', 'mean'), Avg_CR=('Conversion_Rate', 'mean'),
    Campaigns=('Campaign_ID', 'count')
).reset_index().sort_values('Avg_ROI', ascending=False)

add_heading_styled('8.1 Language Performance', level=2)
add_data_table(
    ['Language', 'Campaigns', 'Avg ROI', 'Avg CR'],
    [[row['Language'], f"{row['Campaigns']:,}", f"{row['Avg_ROI']:.2f}", f"{row['Avg_CR']*100:.1f}%"]
     for _, row in lang_stats.iterrows()]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. COST EFFICIENCY & BUDGET OPTIMIZATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('9. Cost Efficiency & Budget Optimization', level=1)

add_body(
    f'Across all campaigns, total acquisition spending reached ${total_spend/1e9:.2f}B, generating '
    f'an estimated {total_conversions/1e6:.1f}M conversions. Understanding the relationship between '
    f'spend and returns is critical for budget optimization.'
)

add_figure('05_cost_efficiency', 'Figure 8: Acquisition Cost vs. ROI scatter plot, color-coded by channel. Each point represents a campaign.')

add_body(
    'The scatter plot reveals that ROI is largely independent of acquisition cost magnitude within '
    'the observed range. This finding suggests that higher spend does not inherently drive higher returns—'
    'rather, creative quality, targeting precision, and channel-campaign fit are the primary determinants '
    'of ROI performance.'
)

# Cost analysis by channel
cost_ch = df.groupby('Channel_Used').agg(
    Avg_Cost=('Acquisition_Cost', 'mean'),
    Avg_CPA=('CPA', 'mean'),
    Avg_ROI=('ROI', 'mean'),
    Total_Spend=('Acquisition_Cost', 'sum')
).reset_index().sort_values('Avg_CPA')

add_heading_styled('9.1 Cost Efficiency by Channel', level=2)
add_data_table(
    ['Channel', 'Avg Acq. Cost', 'Avg CPA', 'Avg ROI', 'Total Spend'],
    [[row['Channel_Used'], f"${row['Avg_Cost']:,.0f}", f"${row['Avg_CPA']:,.0f}",
      f"{row['Avg_ROI']:.2f}", f"${row['Total_Spend']/1e9:.2f}B"]
     for _, row in cost_ch.iterrows()]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 10. TEMPORAL TRENDS & SEASONALITY
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('10. Temporal Trends & Seasonality', level=1)

add_body(
    'Analyzing campaign performance over time reveals seasonal patterns and trends that should inform '
    'campaign scheduling and budget pacing throughout the fiscal year.'
)

add_figure('06_monthly_trends', 'Figure 9: Monthly trends for ROI, Conversion Rate, Total Clicks, and Engagement Score across 2021.')

add_body(
    'Monthly performance metrics remain remarkably stable throughout 2021, with no pronounced seasonal '
    'spikes or troughs. This stability indicates consistent campaign execution and suggests that the '
    'marketing teams maintained disciplined practices year-round. For 2022 planning, budget should be '
    'distributed evenly unless external factors (product launches, seasonal demand) warrant concentration.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 11. COMPANY BENCHMARKING
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('11. Company Benchmarking', level=1)

add_body(
    'Five companies participated in the marketing portfolio. Benchmarking their performance identifies '
    'best practices and areas for cross-organizational learning.'
)

add_figure('08_company_comparison', 'Figure 10: Company-level comparison of ROI, Conversion Rate, and Total Spend.')

comp_sorted = company_stats.sort_values('Avg_ROI', ascending=False)
add_data_table(
    ['Company', 'Avg ROI', 'Avg CR', 'Total Spend', 'Total Conversions', 'Avg Engagement'],
    [[row['Company'], f"{row['Avg_ROI']:.2f}", f"{row['Avg_CR']*100:.1f}%",
      f"${row['Total_Spend']/1e9:.2f}B", f"{row['Total_Conversions']/1e6:.2f}M",
      f"{row['Avg_Engagement']:.2f}"]
     for _, row in comp_sorted.iterrows()]
)

add_heading_styled('11.1 Campaign Duration Analysis', level=2)
add_figure('09_duration_impact', 'Figure 11: Impact of campaign duration on ROI, Conversion Rate, and Engagement Score.')

add_body(
    'Campaign duration (15, 30, 45, or 60 days) shows minimal impact on performance metrics. '
    'This suggests that the current campaign structures achieve their impact early, and longer durations '
    'do not meaningfully improve outcomes. Consider shorter, more frequent campaigns to increase '
    'testing velocity and creative freshness.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 12. STATISTICAL SIGNIFICANCE TESTING
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('12. Statistical Significance Testing (ANOVA)', level=1)

add_body(
    'To determine whether the observed differences across channels, campaign types, audiences, and '
    'other groupings are statistically meaningful — or simply due to random variation — we applied '
    'one-way Analysis of Variance (ANOVA) tests. ANOVA compares group means and returns an F-statistic '
    'and p-value. A p-value below 0.05 indicates a statistically significant difference at the 95% '
    'confidence level.'
)

add_data_table(
    ['Variable Tested', 'F-Statistic', 'p-Value', 'Significant at 95%?'],
    [[var, f"{res['F']:.2f}", f"{res['p']:.4f}" if res['p'] >= 0.0001 else f"{res['p']:.2e}",
      'Yes' if res['significant'] else 'No']
     for var, res in anova_results.items()]
)

# Interpret results
sig_vars = [v for v, r in anova_results.items() if r['significant']]
nonsig_vars = [v for v, r in anova_results.items() if not r['significant']]

if sig_vars:
    add_body(
        f'The ANOVA tests reveal statistically significant differences in ROI across: '
        f'{", ".join(sig_vars)}. This means the differences in average ROI between groups within '
        f'these categories are unlikely to be due to chance alone.'
    )
if nonsig_vars:
    add_body(
        f'Conversely, the following variables showed no statistically significant difference in ROI: '
        f'{", ".join(nonsig_vars)}. The observed variation in average ROI across these groups falls '
        f'within the range expected from random sampling variation.'
    )

add_body(
    'These results are critical for decision-making: investing resources to optimize along statistically '
    'significant dimensions (where real differences exist) will yield returns, while optimizing along '
    'non-significant dimensions may be wasted effort.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 13. PREDICTIVE MODELING & FEATURE IMPORTANCE
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('13. Predictive Modeling & Feature Importance', level=1)

add_body(
    'Beyond descriptive analysis, we trained a machine learning model to predict campaign ROI based on '
    'all available input features. This serves two purposes: (1) quantifying how predictable ROI is from '
    'campaign characteristics, and (2) identifying which features have the most influence on outcomes.'
)

add_heading_styled('13.1 Model Design', level=2)
add_body(
    'We used a Random Forest Regressor — an ensemble of 200 decision trees — trained on 80% of the data '
    'and evaluated on a held-out 20% test set (40,000 campaigns). Random Forest was chosen for its ability '
    'to capture non-linear relationships, handle mixed data types, and provide interpretable feature importance scores.'
)

add_heading_styled('13.2 Model Performance', level=2)

add_kpi_table({
    'R² Score': f"{model_results['r2']:.4f}",
    'Mean Absolute Error': f"{model_results['mae']:.2f}",
    'Test Set Size': '40,000',
    'Trees in Forest': '200'
})

r2 = model_results['r2']
if r2 < 0.05:
    model_interp = (
        f'The model achieved an R² of {r2:.4f}, meaning it explains only {r2*100:.1f}% of the variance in ROI. '
        f'This very low predictive power is itself a finding: it tells us that ROI is largely independent of the '
        f'observable campaign features in this dataset. In other words, the features recorded here — channel, '
        f'campaign type, audience, location, duration — do not meaningfully determine whether a campaign achieves '
        f'high or low ROI. This suggests that unmeasured factors (creative quality, market timing, competitive '
        f'dynamics, ad copy, landing page optimization) are the primary drivers of campaign success.'
    )
elif r2 < 0.3:
    model_interp = (
        f'The model achieved an R² of {r2:.4f}, explaining {r2*100:.1f}% of ROI variance. This moderate '
        f'predictive power suggests that some campaign features contribute to ROI outcomes, but significant '
        f'variance remains unexplained by the available data.'
    )
else:
    model_interp = (
        f'The model achieved an R² of {r2:.4f}, explaining {r2*100:.1f}% of ROI variance. This indicates '
        f'that campaign features collectively have meaningful predictive power for ROI outcomes.'
    )
add_body(model_interp)

add_figure('13_predicted_vs_actual',
    f'Figure 12: Predicted vs. Actual ROI (R² = {r2:.4f}). Points near the red diagonal indicate accurate predictions.')

add_heading_styled('13.3 Feature Importance', level=2)
add_body(
    'Feature importance scores indicate how much each variable contributes to the model\'s predictions. '
    'Higher scores mean the variable has more influence on predicted ROI.'
)

add_figure('12_feature_importance', 'Figure 13: Random Forest feature importance ranking for ROI prediction.')

top3 = model_results['importances'].head(3)
add_body(
    f'The top three predictive features are: {top3.iloc[0]["Feature"]} '
    f'({top3.iloc[0]["Importance"]:.3f}), {top3.iloc[1]["Feature"]} '
    f'({top3.iloc[1]["Importance"]:.3f}), and {top3.iloc[2]["Feature"]} '
    f'({top3.iloc[2]["Importance"]:.3f}). '
    f'This ranking reveals which levers marketers should focus on when designing campaigns.'
)

add_data_table(
    ['Rank', 'Feature', 'Importance Score', 'Cumulative %'],
    [[str(i+1), row['Feature'], f"{row['Importance']:.4f}",
      f"{model_results['importances']['Importance'].iloc[:i+1].sum()*100:.1f}%"]
     for i, (_, row) in enumerate(model_results['importances'].iterrows())]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 14. CORRELATION ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('14. Correlation Analysis', level=1)

add_body(
    'The correlation matrix reveals linear relationships between all numeric campaign metrics. '
    'Values range from -1 (perfect negative correlation) to +1 (perfect positive correlation), '
    'with 0 indicating no linear relationship. Understanding these correlations helps identify '
    'which metrics move together and which are independent.'
)

add_figure('14_correlation_matrix', 'Figure 14: Correlation matrix of all numeric campaign metrics. Blue = positive correlation, Red = negative correlation.')

# Find notable correlations
notable = []
for i in range(len(corr_matrix.columns)):
    for j in range(i+1, len(corr_matrix.columns)):
        val = corr_matrix.iloc[i, j]
        if abs(val) > 0.3:
            notable.append((corr_matrix.columns[i], corr_matrix.columns[j], val))

if notable:
    add_heading_styled('14.1 Notable Correlations', level=2)
    for col1, col2, val in sorted(notable, key=lambda x: abs(x[2]), reverse=True):
        direction = "positive" if val > 0 else "negative"
        strength = "strong" if abs(val) > 0.7 else "moderate"
        doc.add_paragraph(
            f'{col1} and {col2}: {strength} {direction} correlation (r = {val:.2f})',
            style='List Bullet'
        )

add_body(
    'The correlation analysis complements the ANOVA and predictive modeling results by showing '
    'the pairwise linear relationships between metrics. Low correlations between ROI and other '
    'metrics are consistent with the low R² from the predictive model, reinforcing that ROI is '
    'not strongly determined by any single observable metric.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 15. KEY FINDINGS & STRATEGIC RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('15. Key Findings & Strategic Recommendations', level=1)

add_heading_styled('15.1 Key Findings', level=2)

key_findings = [
    ('Balanced Channel Portfolio', 'All six marketing channels deliver comparable ROI, indicating mature optimization across platforms. No single channel dramatically outperforms or underperforms, reducing concentration risk.'),
    ('Social Media Parity', 'Instagram, Facebook, and YouTube perform similarly on aggregate metrics. Platform selection should be driven by campaign objectives (brand awareness vs. conversion) rather than ROI differences alone.'),
    ('Duration Independence', 'Campaign duration has negligible impact on ROI, conversion rate, or engagement. This challenges the assumption that longer campaigns inherently perform better.'),
    ('Uniform Geographic Performance', 'All five metropolitan markets deliver similar returns, suggesting that national strategies are appropriately calibrated without significant regional underperformance.'),
    ('Cost-ROI Decorrelation', 'Higher acquisition costs do not predict higher ROI. Budget efficiency comes from targeting and creative optimization, not increased spend.'),
    ('Statistical Validation', f'ANOVA testing confirmed that {len([v for v,r in anova_results.items() if r["significant"]])} of 6 tested variables show statistically significant differences in ROI at the 95% confidence level, providing rigorous evidence for which optimization levers matter.'),
    ('Low Feature Predictability', f'The Random Forest predictive model achieved an R² of {model_results["r2"]:.4f}, indicating that observable campaign features explain only {model_results["r2"]*100:.1f}% of ROI variance. This strongly suggests that unmeasured factors — creative quality, ad copy, landing pages — are the true drivers of campaign success.'),
]

for title, desc in key_findings:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(6)

add_heading_styled('15.2 Strategic Recommendations', level=2)

recommendations = [
    ('Implement Channel-Campaign Matching', 'Use the interaction heatmap (Figure 3) to match campaign types to their highest-ROI channels rather than distributing uniformly. Even small ROI improvements at this campaign volume translate to significant absolute returns.'),
    ('Shift to Shorter Campaign Cycles', 'Given that duration does not improve performance, adopt 15-30 day campaign windows to increase iteration speed, enable more A/B testing, and maintain creative freshness.'),
    ('Develop Segment-Specific Creative', 'While current creative resonates consistently across segments, targeted messaging for top-performing segments could unlock incremental conversion gains.'),
    ('Cross-Platform Social Sequencing', 'Build awareness on YouTube (video) and Instagram (visual), then retarget engaged users on Facebook for conversion. This funnel-based approach leverages each platform\'s strengths.'),
    ('Invest in Attribution Modeling', 'The uniform performance across channels may mask cross-channel effects. Implement multi-touch attribution to understand how channels work together, not just independently.'),
    ('Regional Creative Testing', 'For markets with marginally lower ROI, run localized A/B tests on creative messaging and channel mix before making wholesale budget reallocation decisions.'),
]

for i, (title, desc) in enumerate(recommendations, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}: ')
    run.bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 16. LIMITATIONS & FUTURE RESEARCH
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('16. Limitations & Future Research', level=1)

add_heading_styled('16.1 Limitations', level=2)
limitations = [
    'Synthetic Data: This dataset is synthetically generated, which may not capture real-world distributions, seasonality effects, or competitive dynamics that influence actual campaign performance.',
    'Single-Touch Attribution: The analysis treats each campaign independently. In practice, customers interact with multiple touchpoints before converting, and cross-channel attribution is essential.',
    'Missing Variables: Key factors such as creative quality scores, bid strategies, audience overlap between campaigns, and competitive intensity are not represented in the data.',
    'Static Audience Definitions: The five audience segments are broad categories. Real-world targeting uses hundreds of micro-segments and lookalike audiences.',
    'No A/B Testing Framework: Without controlled experiments, observed differences cannot be attributed to channel or campaign type with statistical certainty.'
]
for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

add_heading_styled('16.2 Future Research Directions', level=2)
future = [
    'Multi-Touch Attribution Modeling: Implement Shapley value or Markov chain attribution models to quantify cross-channel effects.',
    'Causal Inference: Apply difference-in-differences or regression discontinuity designs to establish causal relationships between channel allocation and outcomes.',
    'Predictive Modeling: Build machine learning models (gradient boosting, neural networks) to predict campaign ROI based on input features, enabling proactive budget optimization.',
    'Customer Lifetime Value Integration: Extend the analysis beyond acquisition cost to include long-term customer value, retention rates, and repeat purchase behavior.',
    'Real-Time Optimization: Develop dashboards for real-time campaign monitoring with automated budget reallocation triggers based on early performance signals.'
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 17. APPENDIX
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled('17. Appendix', level=1)

add_heading_styled('A. Complete Channel Statistics', level=2)

full_ch = df.groupby('Channel_Used').agg(
    Campaigns=('Campaign_ID', 'count'),
    Avg_ROI=('ROI', 'mean'), Std_ROI=('ROI', 'std'),
    Avg_CR=('Conversion_Rate', 'mean'),
    Avg_CTR=('CTR', 'mean'),
    Avg_Engagement=('Engagement_Score', 'mean'),
    Avg_Cost=('Acquisition_Cost', 'mean'),
    Avg_CPA=('CPA', 'mean'),
    Total_Clicks=('Clicks', 'sum'),
    Total_Impressions=('Impressions', 'sum'),
    Total_Conversions=('Est_Conversions', 'sum')
).reset_index()

add_data_table(
    ['Channel', 'N', 'ROI (M)', 'ROI (SD)', 'CR', 'CTR', 'Eng.', 'Avg Cost', 'CPA'],
    [[row['Channel_Used'], f"{row['Campaigns']:,}",
      f"{row['Avg_ROI']:.2f}", f"{row['Std_ROI']:.2f}",
      f"{row['Avg_CR']*100:.1f}%", f"{row['Avg_CTR']*100:.1f}%",
      f"{row['Avg_Engagement']:.2f}", f"${row['Avg_Cost']:,.0f}",
      f"${row['Avg_CPA']:,.0f}"]
     for _, row in full_ch.iterrows()]
)

add_heading_styled('B. Technical Notes', level=2)
add_body('Tools & Libraries: Python 3.14, pandas 3.0, NumPy 2.4, matplotlib 3.10, seaborn 0.13, scikit-learn (Random Forest Regressor), scipy (ANOVA / statistical testing), python-docx 1.2')
add_body('Data Source: Kaggle — Marketing Campaign Performance Dataset (CC0 License, 200,000 records)')
add_body('All figures generated at 150 DPI. Statistical aggregations use arithmetic means unless otherwise noted.')

add_heading_styled('C. Glossary', level=2)
glossary = [
    ('ROI', 'Return on Investment — the ratio of net profit to total investment cost.'),
    ('CTR', 'Click-Through Rate — the percentage of impressions that resulted in a click.'),
    ('CPA', 'Cost Per Acquisition — the average cost to acquire one converting customer.'),
    ('CR', 'Conversion Rate — the percentage of clicks that resulted in a desired action.'),
    ('Engagement Score', 'A 1-10 rating measuring the depth and quality of user interaction with campaign content.'),
]
for term, defn in glossary:
    p = doc.add_paragraph()
    run = p.add_run(f'{term}: ')
    run.bold = True
    p.add_run(defn)

# ── Save ────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_DOCX)
print(f"\nDocument saved: {OUTPUT_DOCX}")
print("Done!")
